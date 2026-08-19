import io
import fitz  # PyMuPDF
import docx
from fastapi.testclient import TestClient
from main import app
from parser.text_cleaner import clean_text
from parser.extractor import extract_candidate_info, extract_phone, extract_urls

client = TestClient(app)


def create_sample_pdf(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def create_two_column_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    # Left column: Personal Details & Skills
    page.insert_text((50, 50), "Yashu Sharma\nEmail: yashu@example.com\nPhone: +91 9876543210")
    page.insert_text((50, 150), "Skills\nPython, FastAPI, React")
    # Right column: Summary & Experience
    page.insert_text((300, 50), "Summary\nSoftware developer building APIs.")
    page.insert_text((300, 150), "Experience\nBackend Lead - Tech Corp\n2022 - Present")
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def create_image_scanned_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    # Draw an image box without text
    pix = fitz.Pixmap(fitz.csRGB, fitz.Rect(0, 0, 100, 100), False)
    page.insert_image(fitz.Rect(50, 50, 250, 250), pixmap=pix)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def create_docx_with_table() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Ananya Sharma\nEmail: ananya@example.com")
    doc.add_paragraph("Education")
    table = doc.add_table(rows=2, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Degree"
    hdr_cells[1].text = "Institution"
    hdr_cells[2].text = "Year"
    row_cells = table.rows[1].cells
    row_cells[0].text = "B.Tech Computer Science"
    row_cells[1].text = "IIT Bombay"
    row_cells[2].text = "2024"
    doc.add_paragraph("Skills\nJava, Python, SQL")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_sample_docx(text: str) -> bytes:
    doc = docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 1. Text Cleaner & Noise Tests ──
def test_text_cleaner():
    raw = "John   Doe\r\n\r\n\r\nEmail: john@example.com\n\n\nSkills:   Python,   React  "
    cleaned = clean_text(raw)
    assert "John Doe" in cleaned
    assert "john@example.com" in cleaned
    assert "Python, React" in cleaned
    assert "\r" not in cleaned


def test_header_footer_noise_cleaning():
    raw = "John Doe\nPage 1 of 2\nSoftware Engineer\nCurriculum Vitae\nPython, Docker"
    cleaned = clean_text(raw)
    assert "Page 1 of 2" not in cleaned
    assert "Curriculum Vitae" not in cleaned
    assert "John Doe" in cleaned
    assert "Software Engineer" in cleaned


# ── 2. Field Extractor & Edge Cases ──
def test_field_extraction():
    sample_text = """John Doe
Software Engineer
Email: john.doe@example.com
Phone: +91 9876543210
Location: Bengaluru, Karnataka
LinkedIn: https://linkedin.com/in/johndoe
GitHub: https://github.com/johndoe
Portfolio: https://johndoe.dev

Summary
Experienced software engineer skilled in building scalable cloud web applications.

Technical Skills
Python, React, FastAPI, PostgreSQL, Docker, Git

Education
Bachelor of Technology in Computer Science
2020 - 2024
CGPA: 8.5/10

Work Experience
Software Engineer - TechCorp Solutions
2022 - Present
• Architected microservices boosting throughput by 40%.
• Built automated CI/CD pipeline with GitHub Actions.

Projects
Distributed Analytics Engine
• Achieved sub-10ms response latency using Python and Redis.

Certifications
AWS Solutions Architect

Achievements
• Winner of National Cloud Hackathon 2023
"""

    cleaned = clean_text(sample_text)
    candidate = extract_candidate_info(cleaned)

    assert candidate.name == "John Doe"
    assert candidate.email == "john.doe@example.com"
    assert candidate.phone == "+91 9876543210"
    assert candidate.linkedin == "https://linkedin.com/in/johndoe"
    assert candidate.github == "https://github.com/johndoe"
    assert candidate.portfolio == "https://johndoe.dev"
    assert candidate.location.city == "Bengaluru"
    assert candidate.location.state == "Karnataka"
    assert candidate.location.country == "India"
    assert "Python" in candidate.skills
    assert "FastAPI" in candidate.skills
    assert "Docker" in candidate.skills
    assert len(candidate.education) > 0
    assert candidate.education[0].degree == "Bachelor of Technology"
    assert len(candidate.experience) > 0
    assert len(candidate.projects) > 0
    assert len(candidate.certifications) > 0
    assert len(candidate.achievements) > 0


def test_unicode_name_and_special_chars():
    sample_text = """आरव शर्मा
Backend Specialist
Email: aarav@example.com
Phone: 9876543210
Location: Delhi, India

Skills
Python, FastAPI, Redis
"""
    cleaned = clean_text(sample_text)
    candidate = extract_candidate_info(cleaned)
    assert candidate.name == "आरव शर्मा"
    assert candidate.email == "aarav@example.com"
    assert "Python" in candidate.skills


def test_multiple_phone_formats():
    phones = ["+91 9876543210", "+91-98765-43210", "(555) 000-0000", "+1.800.555.0199", "9876543210"]
    for p in phones:
        extracted = extract_phone(f"Contact me at {p} anytime.")
        assert extracted is not None


# ── 3. Endpoint & Document Structure Edge Cases ──
def test_health_check():
    response = client.get("/")
    assert response.status_code == 200
    assert response.json()["status"] == "ok"


def test_pdf_upload():
    sample_text = """Yash Pal
Full Stack Developer
yash@example.com
+91 9123456789
Bengaluru, Karnataka
https://linkedin.com/in/yashpal
https://github.com/yashpal

Summary
Full stack software developer.

Technical Skills
Python, TypeScript, React, PostgreSQL, Docker

Education
Bachelor of Engineering
2020 - 2024
"""
    pdf_bytes = create_sample_pdf(sample_text)
    response = client.post(
        "/parse",
        files={"file": ("resume.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["meta"]["file_type"] == "pdf"
    assert data["meta"]["requires_manual_entry"] is False
    assert data["candidate"]["name"] == "Yash Pal"
    assert data["candidate"]["email"] == "yash@example.com"
    assert "Python" in data["candidate"]["skills"]


def test_two_column_pdf():
    pdf_bytes = create_two_column_pdf()
    response = client.post(
        "/parse",
        files={"file": ("twocolumn.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["candidate"]["name"] == "Yashu Sharma"
    assert data["candidate"]["email"] == "yashu@example.com"
    assert "Python" in data["candidate"]["skills"]


def test_docx_upload():
    sample_text = """Ananya Sharma
Backend Developer
ananya@example.com
9876543210
Mumbai, Maharashtra

Technical Skills
Java, Spring Boot, MySQL, Kubernetes
"""
    docx_bytes = create_sample_docx(sample_text)
    response = client.post(
        "/parse",
        files={"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["meta"]["file_type"] == "docx"
    assert data["candidate"]["name"] == "Ananya Sharma"
    assert data["candidate"]["email"] == "ananya@example.com"
    assert "Java" in data["candidate"]["skills"]


def test_docx_with_table():
    docx_bytes = create_docx_with_table()
    response = client.post(
        "/parse",
        files={"file": ("table_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["candidate"]["name"] == "Ananya Sharma"
    assert data["candidate"]["email"] == "ananya@example.com"
    assert "Java" in data["candidate"]["skills"]


def test_scanned_pdf():
    pdf_bytes = create_image_scanned_pdf()
    response = client.post(
        "/parse",
        files={"file": ("scanned.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["meta"]["requires_manual_entry"] is True
    assert data["meta"]["reason"] == "scanned_pdf"


def test_unsupported_file_type():
    response = client.post(
        "/parse",
        files={"file": ("test.txt", b"Some raw text", "text/plain")}
    )
    assert response.status_code == 400
    data = response.json()
    assert data["success"] is False
    assert data["error"]["code"] == "UNSUPPORTED_FILE_TYPE"
